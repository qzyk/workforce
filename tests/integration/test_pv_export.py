"""
Integration tests pentru Faza 14 - PV export DOCX + PDF.
"""

from datetime import date

import pytest


@pytest.fixture
def setup_pv(app, admin_user):
    from models import db, Proiect, Contract, ProcesVerbal
    from services.feature_flags import set_flag
    with app.app_context():
        set_flag('controale-contract', True, commit=True)
        # Cleanup
        ProcesVerbal.query.delete()
        Contract.query.delete()
        Proiect.query.filter_by(cod_proiect='PV-PRJ').delete()
        db.session.commit()

        p = Proiect(cod_proiect='PV-PRJ', nume='PV Test',
                    data_start=date(2026, 1, 1), status='activ')
        db.session.add(p); db.session.commit()
        c = Contract(proiect_id=p.id, nr_contract='PV-CTR',
                     data_semnare=date(2026, 1, 15), status='activ',
                     beneficiar='Beneficiar Test', antreprenor='Antreprenor Test')
        db.session.add(c); db.session.commit()
        pv = ProcesVerbal(
            proiect_id=p.id, contract_id=c.id,
            tip='predare_amplasament', numar='PV-001',
            data_emitere=date(2026, 2, 1),
            obiect='Predarea amplasamentului lucrarii',
            concluzii='Amplasamentul a fost predat fara observatii.',
        )
        pv.participanti = [
            {'nume': 'Ion Popescu', 'functie': 'Diriginte santier',
             'organizatie': 'Beneficiar SRL'},
            {'nume': 'Maria Ionescu', 'functie': 'Sef proiect',
             'organizatie': 'Antreprenor SA'},
        ]
        db.session.add(pv); db.session.commit()
        yield {'pv_id': pv.id, 'proiect_id': p.id, 'contract_id': c.id}
    with app.app_context():
        set_flag('controale-contract', False, commit=True)
        ProcesVerbal.query.delete()
        Contract.query.delete()
        Proiect.query.filter_by(cod_proiect='PV-PRJ').delete()
        db.session.commit()


class TestPVExportDOCX:
    def test_export_docx_endpoint(self, authenticated_client, setup_pv):
        r = authenticated_client.get(f'/contracte/pv/{setup_pv["pv_id"]}/export/docx')
        assert r.status_code == 200
        assert r.headers['Content-Type'].startswith(
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        # DOCX = ZIP, magic bytes PK
        assert r.data[:2] == b'PK'

    def test_docx_contine_titlu_si_participanti(self, app, setup_pv):
        from services.pv_generator import genereaza_pv_docx
        from docx import Document
        with app.app_context():
            path = genereaza_pv_docx(setup_pv['pv_id'])
        # Citesc DOCX-ul generat si verific continutul
        doc = Document(path)
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        # Titlu
        assert 'PREDARE-PRIMIRE AMPLASAMENT' in full_text
        # Numar
        assert 'PV-001' in full_text
        # Concluzii
        assert 'fara observatii' in full_text
        # Participanti (in tabel)
        full_tables = '\n'.join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        assert 'Ion Popescu' in full_tables
        assert 'Maria Ionescu' in full_tables


class TestPVExportPDF:
    def test_export_pdf_endpoint(self, authenticated_client, setup_pv):
        r = authenticated_client.get(f'/contracte/pv/{setup_pv["pv_id"]}/export/pdf')
        assert r.status_code == 200
        assert r.headers['Content-Type'].startswith('application/pdf')
        assert r.data[:4] == b'%PDF'

    def test_pdf_contine_continut(self, app, setup_pv):
        import pypdf
        from services.pv_generator import genereaza_pv_pdf
        with app.app_context():
            path = genereaza_pv_pdf(setup_pv['pv_id'])
        r = pypdf.PdfReader(path)
        all_text = '\n'.join(p.extract_text() for p in r.pages)
        assert 'PREDARE-PRIMIRE AMPLASAMENT' in all_text
        assert 'PV-001' in all_text


class TestPVTipuriDiferite:
    @pytest.mark.parametrize('tip,expected_titlu_fragment', [
        ('predare_amplasament', 'PREDARE-PRIMIRE AMPLASAMENT'),
        ('receptie_proiectare', 'RECEPTIE PROIECTARE'),
        ('receptie_partiala', 'RECEPTIE PARTIALA'),
        ('receptie_finala', 'RECEPTIE FINALA'),
    ])
    def test_titlu_per_tip(self, app, setup_pv, tip, expected_titlu_fragment):
        from models import db, ProcesVerbal
        from services.pv_generator import genereaza_pv_docx
        from docx import Document
        with app.app_context():
            pv = ProcesVerbal.query.get(setup_pv['pv_id'])
            pv.tip = tip
            db.session.commit()
            path = genereaza_pv_docx(pv.id)
        doc = Document(path)
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        assert expected_titlu_fragment in full_text
