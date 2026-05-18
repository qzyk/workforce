"""
Generator Procese Verbale: DOCX (python-docx) + PDF (reportlab) (Faza 14).

4 tipuri PV cu template-uri specifice format romanesc:
  - predare_amplasament: HG907/2016 - predarea-primirea amplasamentului
  - receptie_proiectare: receptia documentatiei tehnice
  - receptie_partiala:   receptia stadiu fizic
  - receptie_finala:     receptia finala lucrari

Toate template-urile sunt cod Python care construieste documentul direct
(NU Jinja runtime in .docx). Avantaj: zero file IO pe template files,
ușor de versionat / modificat / testat.
"""

from __future__ import annotations

import os
from datetime import date, datetime
from typing import Optional

from flask import current_app

from models import db, ProcesVerbal


# ============================================================
# Helpers
# ============================================================

def _upload_dir(subdir: str) -> str:
    base = current_app.config.get(
        'UPLOAD_FOLDER',
        os.path.join(current_app.root_path, 'uploads')
    )
    path = os.path.join(base, subdir)
    os.makedirs(path, exist_ok=True)
    return path


def _titlu_pv(tip: str) -> str:
    return {
        'predare_amplasament': 'PROCES VERBAL DE PREDARE-PRIMIRE AMPLASAMENT',
        'receptie_proiectare': 'PROCES VERBAL DE RECEPTIE PROIECTARE',
        'receptie_partiala':   'PROCES VERBAL DE RECEPTIE PARTIALA (STADIU FIZIC)',
        'receptie_finala':     'PROCES VERBAL DE RECEPTIE FINALA A LUCRARILOR',
        'altul':               'PROCES VERBAL',
    }.get(tip, 'PROCES VERBAL')


def _intro_pv(tip: str) -> str:
    return {
        'predare_amplasament': (
            'Incheiat in conformitate cu prevederile HG 907/2016 si a contractului '
            'de lucrari, prin care beneficiarul preda iar antreprenorul preia '
            'amplasamentul lucrarii in vederea inceperii executiei.'
        ),
        'receptie_proiectare': (
            'Incheiat la finalul fazei de proiectare, prin care beneficiarul receptioneaza '
            'documentatia tehnica elaborata de proiectant.'
        ),
        'receptie_partiala': (
            'Incheiat in baza contractului de lucrari, prin care se receptioneaza '
            'stadiul fizic actual al lucrarilor.'
        ),
        'receptie_finala': (
            'Incheiat la terminarea lucrarilor, in conformitate cu HG 273/1994 modificata, '
            'prin care beneficiarul receptioneaza lucrarile in ansamblu.'
        ),
        'altul': (
            'Incheiat astazi pentru a constata, in baza prezentului proces verbal, '
            'aspectele de mai jos.'
        ),
    }.get(tip, 'Incheiat astazi.')


# ============================================================
# DOCX generator
# ============================================================

def genereaza_pv_docx(pv_id: int) -> str:
    """
    Genereaza un DOCX pentru un ProcesVerbal. Returneaza path-ul fisierului.
    Salveaza si pe modelul `fisier_docx_path`.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    pv = ProcesVerbal.query.get(pv_id)
    if pv is None:
        raise ValueError(f'PV id={pv_id} nu exista.')

    doc = Document()
    # Margins
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # ---- Titlu ----
    titlu = doc.add_paragraph()
    titlu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = titlu.add_run(_titlu_pv(pv.tip))
    run_t.bold = True
    run_t.font.size = Pt(14)
    run_t.font.color.rgb = RGBColor(0x0B, 0x14, 0x26)  # navy obsidian Edifico

    # Numar + data
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        f'Nr. {pv.numar or "____"} din data de {pv.data_emitere}'
    )
    sub_run.italic = True
    sub_run.font.size = Pt(11)

    doc.add_paragraph()  # spatiu

    # ---- Intro ----
    intro = doc.add_paragraph(_intro_pv(pv.tip))
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ---- Date proiect / contract ----
    doc.add_paragraph()
    proiect = pv.proiect
    contract = pv.contract

    p_date = doc.add_paragraph()
    p_date.add_run('Date proiect: ').bold = True
    if proiect:
        p_date.add_run(f'{proiect.cod_proiect} - {proiect.nume}')

    if contract:
        p_ctr = doc.add_paragraph()
        p_ctr.add_run('Contract: ').bold = True
        p_ctr.add_run(f'nr. {contract.nr_contract} din {contract.data_semnare}')
        if contract.beneficiar:
            p_ctr.add_run(f' • Beneficiar: {contract.beneficiar}')
        if contract.antreprenor:
            p_ctr.add_run(f' • Antreprenor: {contract.antreprenor}')

    # ---- Obiect ----
    if pv.obiect:
        doc.add_paragraph()
        h_o = doc.add_paragraph()
        h_o.add_run('Obiectul prezentului PV:').bold = True
        po = doc.add_paragraph(pv.obiect)
        po.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ---- Participanti ----
    if pv.participanti:
        doc.add_paragraph()
        h_p = doc.add_paragraph()
        h_p.add_run('Participanti:').bold = True
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Nume'
        hdr[1].text = 'Functie'
        hdr[2].text = 'Organizatie'
        for part in pv.participanti:
            row = table.add_row().cells
            row[0].text = part.get('nume', '')
            row[1].text = part.get('functie', '')
            row[2].text = part.get('organizatie', '')

    # ---- Concluzii ----
    if pv.concluzii:
        doc.add_paragraph()
        h_c = doc.add_paragraph()
        h_c.add_run('Concluzii:').bold = True
        pc = doc.add_paragraph(pv.concluzii)
        pc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ---- Semnaturi ----
    doc.add_paragraph()
    doc.add_paragraph()
    sem = doc.add_paragraph()
    sem.add_run('Semnaturi:').bold = True
    if pv.participanti:
        for part in pv.participanti:
            sp = doc.add_paragraph()
            sp.add_run(f'{part.get("nume", "")} ({part.get("functie", "")}) ')
            sp.add_run('_______________________')
    else:
        # 2 placeholders generici
        doc.add_paragraph('_______________________   _______________________')

    # Footer
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    foot_run = foot.add_run(
        f'Document generat cu Edifico Workforce - {datetime.utcnow():%Y-%m-%d %H:%M} UTC'
    )
    foot_run.italic = True
    foot_run.font.size = Pt(8)
    foot_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Salveaza
    upload_dir = _upload_dir('procese_verbale')
    filename = f'pv_{pv.tip}_{pv.id}_{datetime.utcnow():%Y%m%d%H%M%S}.docx'
    path = os.path.join(upload_dir, filename)
    doc.save(path)
    pv.fisier_docx_path = path
    db.session.commit()
    return path


# ============================================================
# PDF generator (reportlab)
# ============================================================

def genereaza_pv_pdf(pv_id: int) -> str:
    """Genereaza PDF pentru un ProcesVerbal cu reportlab. Returneaza path."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    )
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_RIGHT

    pv = ProcesVerbal.query.get(pv_id)
    if pv is None:
        raise ValueError(f'PV id={pv_id} nu exista.')

    upload_dir = _upload_dir('procese_verbale')
    filename = f'pv_{pv.tip}_{pv.id}_{datetime.utcnow():%Y%m%d%H%M%S}.pdf'
    path = os.path.join(upload_dir, filename)

    doc = SimpleDocTemplate(path, pagesize=A4,
                            leftMargin=18 * mm, rightMargin=18 * mm,
                            topMargin=18 * mm, bottomMargin=18 * mm)
    styles = getSampleStyleSheet()

    title_st = ParagraphStyle('PVTitle', parent=styles['Title'], fontSize=14,
                              textColor=colors.HexColor('#0B1426'),
                              alignment=TA_CENTER, spaceAfter=6)
    sub_st = ParagraphStyle('PVSub', parent=styles['Normal'], fontSize=10,
                            alignment=TA_CENTER, textColor=colors.grey,
                            spaceAfter=12)
    h_st = ParagraphStyle('PVH', parent=styles['Heading4'], fontSize=11,
                          textColor=colors.HexColor('#0B1426'),
                          spaceBefore=10, spaceAfter=4)
    body_st = ParagraphStyle('PVBody', parent=styles['Normal'], fontSize=10,
                             alignment=TA_JUSTIFY, spaceAfter=6)
    foot_st = ParagraphStyle('PVFoot', parent=styles['Normal'], fontSize=8,
                             alignment=TA_RIGHT, textColor=colors.grey,
                             spaceBefore=20)

    elems = []
    elems.append(Paragraph(_titlu_pv(pv.tip), title_st))
    elems.append(Paragraph(
        f'Nr. {pv.numar or "____"} din data de {pv.data_emitere}', sub_st))

    elems.append(Paragraph(_intro_pv(pv.tip), body_st))

    proiect = pv.proiect
    contract = pv.contract
    if proiect:
        elems.append(Paragraph(
            f'<b>Date proiect:</b> {proiect.cod_proiect} - {proiect.nume}',
            body_st))
    if contract:
        c_info = f'<b>Contract:</b> nr. {contract.nr_contract} din {contract.data_semnare}'
        if contract.beneficiar:
            c_info += f' • <b>Beneficiar:</b> {contract.beneficiar}'
        if contract.antreprenor:
            c_info += f' • <b>Antreprenor:</b> {contract.antreprenor}'
        elems.append(Paragraph(c_info, body_st))

    if pv.obiect:
        elems.append(Paragraph('Obiectul prezentului PV:', h_st))
        elems.append(Paragraph(pv.obiect, body_st))

    if pv.participanti:
        elems.append(Paragraph('Participanti:', h_st))
        rows = [['Nume', 'Functie', 'Organizatie']]
        for p in pv.participanti:
            rows.append([p.get('nume', ''), p.get('functie', ''),
                         p.get('organizatie', '')])
        tbl = Table(rows, colWidths=[60 * mm, 50 * mm, 60 * mm])
        tbl.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#C9A961')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
        ]))
        elems.append(tbl)

    if pv.concluzii:
        elems.append(Paragraph('Concluzii:', h_st))
        elems.append(Paragraph(pv.concluzii, body_st))

    elems.append(Spacer(1, 12 * mm))
    elems.append(Paragraph('Semnaturi:', h_st))
    if pv.participanti:
        for p in pv.participanti:
            nume = p.get('nume', '')
            functie = p.get('functie', '')
            elems.append(Paragraph(
                f'{nume} ({functie}) _______________________', body_st))
    else:
        elems.append(Paragraph(
            '_______________________ &nbsp;&nbsp;&nbsp; _______________________',
            body_st))

    elems.append(Paragraph(
        f'Document generat cu Edifico Workforce - '
        f'{datetime.utcnow():%Y-%m-%d %H:%M} UTC',
        foot_st))

    doc.build(elems)
    pv.fisier_pdf_path = path
    db.session.commit()
    return path
